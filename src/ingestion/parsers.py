"""Document parsers for different file types."""
import mimetypes
from pathlib import Path
from typing import Optional, Callable, Dict
from dataclasses import dataclass

import pypdf
from docx import Document
import pandas as pd
from src.config.logging_config import get_logger

logger = get_logger(__name__)
ParserFunc = Callable[[Path], "ParsedDocument"]

@dataclass
class ParsedDocument:
    """Parsed document with content and metadata."""

    content: str
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    extra_metadata: Optional[dict] = None


class DocumentParser:
    """Base document parser."""

    @staticmethod
    def get_mime_type(file_path: Path) -> str:
        """Get MIME type of file."""
        mime_type, _ = mimetypes.guess_type(str(file_path))
        return mime_type or "application/octet-stream"

    @staticmethod
    def count_words(text: str) -> int:
        """Count words in text."""
        return len(text.split())

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse document based on file type."""
        suffix = file_path.suffix.lower()

        parsers: Dict[str, ParserFunc] = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".txt": self._parse_text,
            ".md": self._parse_text,
            ".csv": self._parse_csv,
            ".xlsx": self._parse_excel,
            ".json": self._parse_json,
        }

        parser_func = parsers.get(suffix)
        if not parser_func:
            raise ValueError(f"Unsupported file type: {suffix}")

        try:
            return parser_func(file_path)
        except Exception as e:
            logger.error("parse_failed", file_path=str(file_path), error=str(e))
            raise

    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """Parse PDF file."""
        logger.info("parsing_pdf", file_path=str(file_path))

        with open(file_path, "rb") as f:
            pdf_reader = pypdf.PdfReader(f)
            page_count = len(pdf_reader.pages)

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    # Add page marker for better chunking
                    text_parts.append(f"\n--- Page {page_num} ---\n{text}")

            content = "\n".join(text_parts)

            # Extract metadata
            metadata = pdf_reader.metadata or {}
            extra_metadata = {
                "author": metadata.get("/Author", ""),
                "title": metadata.get("/Title", ""),
                "subject": metadata.get("/Subject", ""),
            }

        return ParsedDocument(
            content=content,
            page_count=page_count,
            word_count=self.count_words(content),
            extra_metadata=extra_metadata,
        )

    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """Parse DOCX file."""
        logger.info("parsing_docx", file_path=str(file_path))

        doc = Document(str(file_path))

        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        content = "\n\n".join(paragraphs)

        # Extract tables
        tables_content = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text.strip() for cell in row.cells])
            if table_data:
                tables_content.append("\n".join(["\t".join(row) for row in table_data]))

        if tables_content:
            content += "\n\n--- Tables ---\n" + "\n\n".join(tables_content)

        # Core properties
        extra_metadata = {
            "author": doc.core_properties.author or "",
            "title": doc.core_properties.title or "",
        }

        return ParsedDocument(
            content=content,
            page_count=None,
            word_count=self.count_words(content),
            extra_metadata=extra_metadata,
        )

    def _parse_text(self, file_path: Path) -> ParsedDocument:
        """Parse plain text file."""
        logger.info("parsing_text", file_path=str(file_path))

        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        return ParsedDocument(
            content=content,
            page_count=None,
            word_count=self.count_words(content),
        )

    def _parse_csv(self, file_path: Path) -> ParsedDocument:
        """Parse CSV file."""
        logger.info("parsing_csv", file_path=str(file_path))

        df = pd.read_csv(file_path)

        # Convert to formatted text with headers
        content_parts = [
            f"CSV File: {file_path.name}",
            f"Columns: {', '.join(df.columns.tolist())}",
            f"Total Rows: {len(df)}",
            "\n--- Data Sample (first 100 rows) ---\n",
            df.head(100).to_string(index=False),
        ]

        content = "\n".join(content_parts)

        extra_metadata = {
            "columns": df.columns.tolist(),
            "row_count": len(df),
            "column_count": len(df.columns),
        }

        return ParsedDocument(
            content=content,
            page_count=None,
            word_count=self.count_words(content),
            extra_metadata=extra_metadata,
        )

    def _parse_excel(self, file_path: Path) -> ParsedDocument:
        """Parse Excel file."""
        logger.info("parsing_excel", file_path=str(file_path))

        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        sheets_content = []

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            sheet_text = [
                f"\n--- Sheet: {sheet_name} ---",
                f"Columns: {', '.join(df.columns.tolist())}",
                f"Rows: {len(df)}",
                df.head(100).to_string(index=False),
            ]
            sheets_content.append("\n".join(sheet_text))

        content = "\n\n".join(sheets_content)

        extra_metadata = {
            "sheets": excel_file.sheet_names,
            "sheet_count": len(excel_file.sheet_names),
        }

        return ParsedDocument(
            content=content,
            page_count=None,
            word_count=self.count_words(content),
            extra_metadata=extra_metadata,
        )

    def _parse_json(self, file_path: Path) -> ParsedDocument:
        """Parse JSON file."""
        logger.info("parsing_json", file_path=str(file_path))

        import json

        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # Convert to readable text
        content = json.dumps(data, indent=2)

        return ParsedDocument(
            content=content,
            page_count=None,
            word_count=self.count_words(content),
            extra_metadata={"json_keys": list(data.keys()) if isinstance(data, dict) else None},
        )